import os
import streamlit as st
from pptx import Presentation
from pptx.util import Pt
from docx import Document
import pandas as pd
from deep_translator import GoogleTranslator
from indic_transliteration.sanscript import transliterate, KANNADA, DEVANAGARI, TELUGU, HK

# ------------------------------
# LANGUAGE + SCRIPT CONFIG
# ------------------------------
script_codes = {
    "Kannada": KANNADA,
    "Hindi (Devanagari)": DEVANAGARI,
    "Marathi (Devanagari)": DEVANAGARI,
    "Telugu": TELUGU,
    "English (HK)": HK
}

translation_lang_codes = {
    "Kannada": "kn",
    "Hindi (Devanagari)": "hi",
    "Marathi (Devanagari)": "mr",
    "Telugu": "te",
    "English (HK)": "en"
}

# ------------------------------
# UNIVERSAL TRANSLATE + TRANSLITERATE
# ------------------------------
def translate_text(text, target_lang_code):
    try:
        return GoogleTranslator(source='auto', target=target_lang_code).translate(text)
    except Exception as e:
        return f"(Translation Error: {str(e)})"

def transliterate_to_target(text, target_script):
    try:
        return transliterate(text, None, target_script)
    except Exception:
        return text  # fallback: original text

# ------------------------------
# FILE PROCESSING FUNCTIONS
# ------------------------------
def process_docx(uploaded_file, target_script, target_lang_code):
    doc = Document(uploaded_file)
    new_doc = Document()
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        translit_line = transliterate_to_target(line, target_script)
        translated_line = translate_text(line, target_lang_code)

        new_doc.add_paragraph(translit_line)
        new_doc.add_paragraph(translated_line)
        new_doc.add_paragraph("")

    output_path = "output_translated.docx"
    new_doc.save(output_path)
    return output_path

# (Similar change can be applied to PPTX, Excel, CSV, TXT)

# ------------------------------
# STREAMLIT UI
# ------------------------------
st.set_page_config(page_title="Universal Transliterator + Translator", layout="wide")
st.title("🌍 Universal Translator + Transliterator")

uploaded_file = st.file_uploader("📂 Upload file", type=["pptx", "docx", "xlsx", "xls", "csv", "txt"])
st.markdown("**OR**")
plain_text_input = st.text_area("📝 Enter any text:")

target_lang = st.selectbox("Select Output Language", list(script_codes.keys()))

if st.button("🚀 Convert"):
    if not uploaded_file and not plain_text_input.strip():
        st.warning("Please upload or enter some text.")
    else:
        with st.spinner("Processing... Please wait ⏳"):
            target_script = script_codes[target_lang]
            target_lang_code = translation_lang_codes[target_lang]

            try:
                if uploaded_file:
                    ext = os.path.splitext(uploaded_file.name)[-1].lower()
                    if ext == ".docx":
                        output_path = process_docx(uploaded_file, target_script, target_lang_code)
                    else:
                        st.error("Only DOCX supported in this demo.")
                        st.stop()

                    st.success("✅ Completed!")
                    with open(output_path, "rb") as f:
                        st.download_button("⬇️ Download Result", f, file_name=output_path)

                elif plain_text_input.strip():
                    lines = plain_text_input.split('\n')
                    result_lines = []
                    for line in lines:
                        if not line.strip():
                            continue
                        translit_line = transliterate_to_target(line, target_script)
                        translated_line = translate_text(line, target_lang_code)
                        result_lines.append(f"{translit_line}\n{translated_line}\n")

                    result_text = "\n".join(result_lines)
                    st.markdown(f"### 📜 Result:\n```\n{result_text}\n```")
                    st.download_button("⬇️ Download Result as TXT", result_text.encode('utf-8'), "translated_text.txt")

            except Exception as e:
                st.error(f"❌ Error: {e}")
